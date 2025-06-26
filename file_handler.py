import streamlit as st
import PyPDF2
from docx import Document
import pandas as pd
import io

def handle_file_upload(uploaded_file):
    """Handle file upload and processing with LangChain analysis"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'txt':
            return handle_text_file(uploaded_file)
        elif file_extension == 'pdf':
            return handle_pdf_file(uploaded_file)
        elif file_extension == 'docx':
            return handle_docx_file(uploaded_file)
        elif file_extension == 'csv':
            return handle_csv_file(uploaded_file)
        else:
            return f"❌ Unsupported file type: {file_extension}"
            
    except Exception as e:
        return f"❌ Error processing file: {str(e)}"

def handle_text_file(uploaded_file):
    """Handle .txt file upload with LangChain analysis"""
    try:
        content = uploaded_file.read().decode('utf-8')
        word_count = len(content.split())
        
        # Get AI assistant for analysis
        import streamlit as st
        if hasattr(st.session_state, 'ai_assistant'):
            analysis = st.session_state.ai_assistant.analyze_document(content, 'txt')
            return analysis
        
        return f"""📄 **Text File Processed Successfully!** ✅

**File:** {uploaded_file.name}
**Size:** {len(content)} characters
**Words:** {word_count}

**Preview:** 
{content[:300]}{'...' if len(content) > 300 else ''}

🤖 **LangChain Analysis:** Ready for intelligent document analysis! Ask me specific questions about this content, and I'll provide detailed insights using advanced language models.

What would you like to know about this document? 🚀"""

    except Exception as e:
        return f"❌ Error reading text file: {str(e)}"

def handle_pdf_file(uploaded_file):
    """Handle .pdf file upload with LangChain analysis"""
    try:
        # Read PDF content
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text_content = ""
        
        for page in pdf_reader.pages:
            text_content += page.extract_text() + "\n"
        
        word_count = len(text_content.split())
        page_count = len(pdf_reader.pages)
        
        # Get AI assistant for analysis
        import streamlit as st
        if hasattr(st.session_state, 'ai_assistant'):
            analysis = st.session_state.ai_assistant.analyze_document(text_content, 'pdf')
            return analysis
        
        return f"""📄 **PDF Document Processed Successfully!** ✅

**File:** {uploaded_file.name}
**Pages:** {page_count}
**Words:** {word_count}
**Characters:** {len(text_content)}

**Content Preview:**
{text_content[:400]}{'...' if len(text_content) > 400 else ''}

🤖 **LangChain + Groq Analysis:** Your PDF is ready for advanced AI analysis! I can now provide intelligent insights, answer complex questions, and perform deep content analysis.

What would you like to explore from this document? 🚀"""

    except Exception as e:
        return f"❌ Error processing PDF: {str(e)}"

def handle_docx_file(uploaded_file):
    """Handle .docx file upload"""
    try:
        doc = Document(uploaded_file)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        word_count = len(text_content.split())
        paragraph_count = len(doc.paragraphs)
        
        return f"""📄 **Word Document Analyzed Successfully!** ✅

**File:** {uploaded_file.name}
**Paragraphs:** {paragraph_count}
**Words:** {word_count}
**Characters:** {len(text_content)}

**Content Preview:**
{text_content[:400]}{'...' if len(text_content) > 400 else ''}

Your Word document is ready for analysis! I can help you:
• Summarize the main points 📋
• Answer specific questions about content 💡
• Identify key themes and concepts 🔑
• Extract important information 📊

How can I help you with this document? 🤖"""

    except Exception as e:
        return f"❌ Error processing Word document: {str(e)}"

def handle_csv_file(uploaded_file):
    """Handle .csv file upload"""
    try:
        df = pd.read_csv(uploaded_file)
        
        # Basic statistics
        rows, cols = df.shape
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        text_cols = df.select_dtypes(include=['object']).columns.tolist()
        
        # Sample data preview
        preview = df.head(3).to_string(index=False)
        
        return f"""📊 **CSV Data Analyzed Successfully!** ✅

**File:** {uploaded_file.name}
**Dimensions:** {rows} rows × {cols} columns
**Numeric Columns:** {len(numeric_cols)} ({', '.join(numeric_cols[:3])}{'...' if len(numeric_cols) > 3 else ''})
**Text Columns:** {len(text_cols)} ({', '.join(text_cols[:3])}{'...' if len(text_cols) > 3 else ''})

**Data Preview:**
```
{preview}
```

Your CSV data is loaded! I can help you:
• Analyze data patterns and trends 📈
• Generate summaries and statistics 📊
• Answer questions about specific data 🔍
• Identify insights and correlations 💡
• Create visualizations (conceptually) 🎨

What insights are you looking for in this data? 🚀"""

    except Exception as e:
        return f"❌ Error processing CSV file: {str(e)}"

def get_file_summary(file_content, file_type):
    """Generate a summary of the uploaded file"""
    word_count = len(file_content.split()) if isinstance(file_content, str) else 0
    
    summary = f"""
🔍 **File Analysis Summary**

**Type:** {file_type.upper()}
**Content Length:** {len(str(file_content))} characters
**Estimated Words:** {word_count}
**Processing Status:** ✅ Complete

The file has been successfully processed and is ready for analysis!
"""
    
    return summary